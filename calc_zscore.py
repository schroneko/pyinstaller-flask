import datetime
import os
import pathlib

import matplotlib
import matplotlib.pyplot as plt

# import mne
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from scipy import signal
from scipy.integrate import simps

from topograph import get_psds_alpha, get_psds_beta, get_psds_theta, plot_topomap

matplotlib.use("Agg")


def calc_zscore(input_data, input_name):
    input_data = os.path.join(os.getcwd(), "out", input_data)
    out_dir = os.path.join(os.getcwd(), "out")

    if input_data.endswith(".m00"):
        df = pd.read_table(
            input_data,
            delim_whitespace=True,
            skiprows=1,
        )
    # elif input_data.endswith(".edf"):
    #     edf_data = mne.io.read_raw_edf(input_data)
    #     labels = edf_data.ch_names
    #     edf_data = edf_data.get_data()
    #     df = pd.DataFrame(edf_data, index=labels).T
    #     df[:][:] *= 1000000

    if df.columns.str.contains("Fp1-A1").any():
        df = df.filter(
            items=[
                "Fp1-A1",
                "Fp2-A2",
                "C3-A1",
                "C4-A2",
                "O1-A1",
                "O2-A2",
                "T3-A1",
                "T4-A2",
                "F7-A1",
                "F8-A2",
                "T5-A1",
                "T6-A2",
            ]
        )

    # if file extension is .edf
    elif df.columns.str.contains("FP1-REF").any():
        df = df.filter(
            items=[
                "EEG FP1-REF",
                "EEG FP2-REF",
                "EEG C3-REF",
                "EEG C4-REF",
                "EEG O1-REF",
                "EEG O2-REF",
                "EEG T3-REF",
                "EEG T4-REF",
                "EEG F7-REF",
                "EEG F8-REF",
                "EEG T5-REF",
                "EEG T6-REF",
            ]
        )

    df.columns = [
        "Fp1",
        "Fp2",
        "C3",
        "C4",
        "O1",
        "O2",
        "T3",
        "T4",
        "F7",
        "F8",
        "T5",
        "T6",
    ]

    # 欠損値が１つでもある行を削除する
    df = df.dropna(how="any")

    # topomapを作成して保存する
    df_copy = df.T
    df_np = df_copy.values

    # シータ波のスペクトルを求める
    pwrs_theta, _ = get_psds_theta(df_np)

    # アルファ波のスペクトルを求める
    pwrs_alpha, _ = get_psds_alpha(df_np)

    # ベータ波のスペクトルを求める
    pwrs_beta, _ = get_psds_beta(df_np)

    # Zスコアを求める
    eeg_list = [
        "Fp1",
        "Fp2",
        "C3",
        "C4",
        "O1",
        "O2",
        "T3",
        "T4",
        "F7",
        "F8",
        "T5",
        "T6",
    ]

    npy_dir = os.path.join(os.getcwd(), "npy")
    np_load_dataset = [
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_fp1_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_fp1_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_fp1_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_fp2_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_fp2_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_fp2_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_c3_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_c3_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_c3_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_c4_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_c4_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_c4_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_O1_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_O1_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_O1_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_O2_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_O2_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_O2_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_t3_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_t3_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_t3_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_t4_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_t4_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_t4_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_f7_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_f7_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_f7_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_f8_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_f8_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_f8_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_t5_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_t5_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_t5_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_t6_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_t6_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_t6_eval+train.npy"),
        ],
    ]

    j_list = ["アルファ", "ベータ", "シータ"]

    # result_listは結果を文章で出力するために用意したリスト
    result_list = []

    # ３行１６列のリストを作る。そしてアルファ波のtopomapならlist[0][l]をlでイテレートして表示する。numpy配列にしないといけないはず。
    # →そのやり方が分からなかったので３種類のリストを作る！
    result_alpha = []
    result_beta = []
    result_theta = []

    sample_spectrum = []

    alpha_fp1, alpha_fp2 = 0, 0
    rel_alpha_fp1, rel_alpha_fp2 = 0, 0
    beta_theta_fp1, beta_theta_fp2 = 0, 0

    # ある電極での相対スペクトル密度のZスコアを求める
    for i in range(len(eeg_list)):
        if input_data.endswith(".m00"):
            dt = 2 * 10 ** (-3)
        if input_data.endswith(".edf"):
            dt = 4 * 10 ** (-3)
        sf = 1 / dt
        df_analyze = df.iloc[:, i]
        df_analyze_np = df_analyze.values

        # Welch's periodogramを求める
        win = 4 * sf
        freqs, psd = signal.welch(df_analyze_np, sf, nperseg=win)
        freq_res = freqs[1] - freqs[0]

        # ベータ波の相対スペクトル密度を求める
        idx_beta = np.logical_and(freqs >= 12, freqs <= 30)
        beta_power = simps(psd[idx_beta], dx=freq_res)

        # アルファ波の相対スペクトル密度を求める
        idx_alpha = np.logical_and(freqs >= 8, freqs <= 12)
        alpha_power = simps(psd[idx_alpha], dx=freq_res)

        # シータ波の相対スペクトル密度を求める
        idx_theta = np.logical_and(freqs >= 4, freqs <= 8)
        theta_power = simps(psd[idx_theta], dx=freq_res)

        total_power = theta_power + alpha_power + beta_power
        relative_alpha = alpha_power / total_power
        relative_beta = beta_power / total_power
        relative_theta = theta_power / total_power
        relative_list = [relative_alpha, relative_beta, relative_theta]

        # 脳波データの周波数スペクトルを求める
        x = freqs
        y = psd

        plt.plot(x, y)
        plt.ticklabel_format(style="plain", axis="y", useOffset=False)
        plt.title("EEG-" + eeg_list[i])
        plt.xlabel("Frequency [Hz]")
        plt.ylabel("Power Spectral Density [V^2/Hz]")
        plt.xlim(4, 30)  # スペクトルの表示範囲をシータ波帯域〜ベータ波帯域に指定
        plt.savefig(os.path.join(out_dir, "EEG-" + eeg_list[i] + ".png"))
        plt.clf()

        # それぞれの電極で計算する
        for j in range(3):
            # 標準偏差を求める
            sample_spectrum = np.load(np_load_dataset[i][j])
            data_mean = np.mean(sample_spectrum)
            data_std = np.std(sample_spectrum)

            z = (relative_list[j] - data_mean) / data_std
            result_list.append(
                "{}電極の{}波の相対パワースペクトルのZスコアは{}です。".format(
                    eeg_list[i], j_list[j], round(z, 2)
                )
            )

        z1 = (relative_list[0] - data_mean) / data_std
        result_alpha.append(z1)

        z2 = (relative_list[1] - data_mean) / data_std
        result_beta.append(z2)

        z3 = (relative_list[2] - data_mean) / data_std
        result_theta.append(z3)

        pwrs_rel_alpha_fp1 = np_load_dataset[0][0]
        pwrs_rel_beta_fp1 = np_load_dataset[0][1]
        pwrs_rel_theta_fp1 = np_load_dataset[0][2]
        pwrs_rel_alpha_fp2 = np_load_dataset[1][0]
        pwrs_rel_beta_fp2 = np_load_dataset[1][1]
        pwrs_rel_theta_fp2 = np_load_dataset[1][2]

        # Fp1とFp2のアルファ波パワーの差とZスコアを求める
        # Fp1とFp2のシータ波パワー/ベータ波パワーとZスコアを求める
        if eeg_list[i] == "Fp1":
            alpha_fp1 = alpha_power
            rel_alpha_fp1 = relative_alpha
            beta_theta_fp1 = beta_power / theta_power  # = rel_beta / rel_theta

            # beta_thetaの標準偏差を求める
            spectrum_beta_theta_fp1 = np.load(pwrs_rel_beta_fp1) / np.load(
                pwrs_rel_theta_fp1
            )
            z_beta_theta_fp1 = (
                beta_theta_fp1 - np.mean(spectrum_beta_theta_fp1)
            ) / np.std(spectrum_beta_theta_fp1)

        elif eeg_list[i] == "Fp2":
            alpha_fp2 = alpha_power
            rel_alpha_fp2 = relative_alpha
            beta_theta_fp2 = beta_power / theta_power  # = rel_beta / rel_theta

            # rel_alpha_diffの標準偏差を求める
            rel_alpha_diff = np.log(rel_alpha_fp2) - np.log(rel_alpha_fp1)
            spectrum_alpha_diff = np.log(np.load(pwrs_rel_alpha_fp2)) - np.log(
                np.load(pwrs_rel_alpha_fp1)
            )
            z_alpha_diff = (rel_alpha_diff - np.mean(spectrum_alpha_diff)) / np.std(
                spectrum_alpha_diff
            )

            # Calculate the difference between the alpha power Fp1 with Fp2
            alpha_diff = np.log(alpha_fp2) - np.log(alpha_fp1)

            # beta_thetaの標準偏差を求める
            spectrum_beta_theta_fp2 = np.load(pwrs_rel_beta_fp2) / np.load(
                pwrs_rel_theta_fp2
            )
            z_beta_theta_fp2 = (
                beta_theta_fp2 - np.mean(spectrum_beta_theta_fp2)
            ) / np.std(spectrum_beta_theta_fp2)

            # beta_thetaの前頭葉左右差とZスコアを求める
            beta_theta_diff = beta_theta_fp2 - beta_theta_fp1

            spectrum_beta_theta_diff = spectrum_beta_theta_fp2 - spectrum_beta_theta_fp1

            z_beta_theta_diff = (
                beta_theta_diff - np.mean(spectrum_beta_theta_diff)
            ) / np.std(spectrum_beta_theta_diff)

    # シータ波のtopomapを出力する
    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_theta, ax, fig)
    plt.title("theta_topomap")
    # plt.title("シータ波のトポグラフィー")
    plt.savefig(os.path.join(out_dir, "theta_save_topomap.png"))

    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_alpha, ax, fig)
    plt.title("alpha_topomap")
    # plt.title("アルファ波のトポグラフィー")
    plt.savefig(os.path.join(out_dir, "alpha_save_topomap.png"))

    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_beta, ax, fig)
    plt.title("beta_topomap")
    # plt.title("ベータ波のトポグラフィー")
    plt.savefig(os.path.join(out_dir, "beta_save_topomap.png"))

    plt.clf()

    dt_now = datetime.datetime.now()
    document = Document()
    document.add_heading("脳波の解析結果")
    document.add_paragraph(" ")
    document.add_paragraph("作成日：" + dt_now.strftime("%Y年%m月%d日"))

    file_name = os.path.basename(input_data)

    document.add_paragraph("ファイル名：" + file_name)
    document.add_paragraph("インプット名：" + input_name)
    document.add_paragraph(" ")
    document.add_heading("脳波トポグラフィー", level=2)
    document.add_paragraph(
        "各時刻、各電極計12箇所における脳波を周波数毎（シータ波・アルファ波・ベータ波として分類）に分析し、脳波トポグラフィーを描画します。"
    )
    # document.add_paragraph("シータ波の分布は以下のようになります。")
    document.add_heading("シータ波の分布", level=3)
    document.add_picture(
        os.path.join(out_dir, "theta_save_topomap.png"), width=Inches(3.5)
    )
    document.add_page_break()

    # document.add_paragraph("アルファ波の分布は以下のようになります。")
    document.add_heading("アルファ波の分布", level=3)
    document.add_picture(
        os.path.join(out_dir, "alpha_save_topomap.png"), width=Inches(3.5)
    )
    document.add_heading("ベータ波の分布", level=3)
    # document.add_paragraph("ベータ波の分布は以下のようになります。")
    document.add_picture(
        os.path.join(out_dir, "beta_save_topomap.png"), width=Inches(3.5)
    )
    # document.add_paragraph(" ")
    document.add_page_break()

    document.add_heading("Zスコア（標準得点）の算出", level=2)
    document.add_paragraph("正常被験者1297人の脳波データから入力ファイルのZスコアを算出します。")
    document.add_paragraph(
        "Zスコアは、入力データと正常なデータ群の平均値との差を正常なデータ群の標準偏差で除したもので、入力データがどれだけ正常なデータから外れているかを示す指標の一つです。"
    )
    document.add_paragraph("例えば、Z値が+1であるならデータセットの分布から標準偏差分だけ外れている事を示します。")

    for k in range(len(result_list)):
        # document.add_paragraph(" ")
        # document.paragraphs[13 + k].add_run(result_list[k])
        document.add_paragraph(result_list[k])
        if (k + 1) % 3 == 0:
            document.add_paragraph(" ")

    document.add_heading("前頭葉の左右差の分析", level=2)
    document.add_paragraph(
        "前頭葉のアルファ波の左右差やそのZスコア、前頭葉のベータ波とシータ波の比率やそのZスコアが様々な症例と相関性のある事が近年の論文によって示されています。"
    )

    document.add_paragraph("前頭葉のアルファ波左右差の値は" + str("{:.3g}".format(alpha_diff)) + "です。")
    document.add_paragraph(
        "前頭葉のアルファ波左右差の相対パワースペクトルのZスコアは" + str("{:.3g}".format(z_alpha_diff)) + "です。"
    )
    document.add_paragraph(
        "前頭葉（右）のベータ波/シータ波の比率は" + str("{:.3g}".format(beta_theta_fp1)) + "です。"
    )
    document.add_paragraph(
        "前頭葉（左）のベータ波/シータ波の比率は" + str("{:.3g}".format(beta_theta_fp2)) + "です。"
    )
    document.add_paragraph(
        "前頭葉（右）のベータ波/シータ波のZスコアは" + str("{:.3g}".format(z_beta_theta_fp1)) + "です。"
    )
    document.add_paragraph(
        "前頭葉（左）のベータ波/シータ波のZスコアは" + str("{:.3g}".format(z_beta_theta_fp2)) + "です。"
    )
    document.add_paragraph(
        "前頭葉のベータ波/シータ波の左右差は" + str("{:.3g}".format(beta_theta_diff)) + "です。"
    )
    document.add_paragraph(
        "前頭葉のベータ波/シータ波の左右差のZスコアは" + str("{:.3g}".format(z_beta_theta_diff)) + "です。"
    )
    document.add_page_break()

    document.add_heading("パワースペクトル密度の周波数依存性", level=2)
    document.add_paragraph("周波数解析した各電極におけるパワースペクトル密度の周波数依存性を描画します。")
    document.add_paragraph(
        "シータ波帯域が4~8 [Hz]、アルファ波帯域が8~12 [Hz]、ベータ波帯域が12~30 [Hz]となっています。"
    )

    for i in range(len(eeg_list)):
        document.add_picture(
            os.path.join(out_dir, "EEG-" + eeg_list[i] + ".png"), width=Inches(3.5)
        )

    extension = pathlib.PurePath(input_data).suffix

    save_dir = input_data.replace(extension, ".docx")

    document.save(save_dir)
